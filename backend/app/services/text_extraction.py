"""Port of the PDF/DOCX/plaintext extraction logic embedded in
app/api/documents/upload/route.ts's processDocument() — pdf-parse/mammoth
replaced with pypdf/python-docx, same behavior (extract, error on empty text).
"""
import io

import docx
from pypdf import PdfReader


class ExtractionError(Exception):
    pass


def extract_text(data: bytes, extension: str) -> str:
    ext = extension.lower()

    if ext == "pdf":
        reader = PdfReader(io.BytesIO(data))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        if not text.strip():
            raise ExtractionError("Could not extract text from PDF. The file may be scanned or image-only.")
    elif ext == "docx":
        document = docx.Document(io.BytesIO(data))
        text = "\n".join(p.text for p in document.paragraphs)
        if not text.strip():
            raise ExtractionError("Could not extract text from DOCX file. The file may be empty or corrupted.")
    else:
        text = data.decode("utf-8", errors="replace")

    text = text.replace("\x00", "")
    if not text.strip():
        raise ExtractionError("Document contains no extractable text.")

    return text
